import io

from PIL import Image
from PyPDF2 import PdfReader
from docx import Document

from backend.api.services.pdf_service import convert_word_to_pdf


def _build_docx_with_content() -> io.BytesIO:
    document = Document()

    heading = document.add_heading('Quarterly Review', level=1)
    heading.add_run(' 2026')

    paragraph = document.add_paragraph()
    paragraph.add_run('Bold ').bold = True
    paragraph.add_run('Italic ').italic = True
    paragraph.add_run('https://example.com')

    document.add_paragraph('First bullet item', style='List Bullet')
    document.add_paragraph('Second bullet item', style='List Bullet')

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Status'
    table.cell(1, 1).text = 'Ready'

    image = Image.new('RGB', (160, 90), 'white')
    image_buffer = io.BytesIO()
    image.save(image_buffer, format='PNG')
    image_buffer.seek(0)
    document.add_picture(image_buffer)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def test_convert_word_to_pdf_preserves_core_content():
    output = convert_word_to_pdf(_build_docx_with_content())
    assert output.getbuffer().nbytes > 0

    reader = PdfReader(output)
    extracted = '\n'.join((page.extract_text() or '') for page in reader.pages)

    assert 'Quarterly Review' in extracted
    assert 'Bold' in extracted
    assert 'Italic' in extracted
    assert 'https://example.com' in extracted
    assert 'First bullet item' in extracted
    assert 'Second bullet item' in extracted
    assert 'Name' in extracted
    assert 'Ready' in extracted


def test_convert_word_to_pdf_handles_image_only_documents():
    document = Document()
    image = Image.new('RGB', (80, 80), 'navy')
    image_buffer = io.BytesIO()
    image.save(image_buffer, format='PNG')
    image_buffer.seek(0)
    document.add_picture(image_buffer)

    output_buffer = io.BytesIO()
    document.save(output_buffer)
    output_buffer.seek(0)

    output = convert_word_to_pdf(output_buffer)
    assert output.getbuffer().nbytes > 0
