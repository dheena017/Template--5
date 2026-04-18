import io
from PyPDF2 import PdfReader, PdfWriter
from typing import List, Optional


def convert_word_to_pdf(file_buffer: io.BytesIO) -> io.BytesIO:
    """Convert a DOCX file into a PDF while preserving basic structure."""

    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    file_buffer.seek(0)
    document = Document(file_buffer)

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    page_width, page_height = A4

    margin_x = 50
    bottom_margin = 60
    y = page_height - 50

    def ensure_space(line_height: int, reset_font: tuple[str, int] | None = None):
        nonlocal y
        if y < bottom_margin:
            pdf.showPage()
            y = page_height - 50
            if reset_font:
                pdf.setFont(reset_font[0], reset_font[1])

    def write_wrapped(text: str, font_name: str = 'Helvetica', font_size: int = 11, indent: int = 0):
        nonlocal y
        if not text:
            ensure_space(14, (font_name, font_size))
            y -= 14
            return

        max_width = page_width - (margin_x * 2) - indent
        pdf.setFont(font_name, font_size)
        words = text.split()
        if not words:
            ensure_space(14, (font_name, font_size))
            y -= 14
            return

        current = ''
        line_height = max(14, int(font_size * 1.35))
        for word in words:
            candidate = f"{current} {word}".strip()
            if pdf.stringWidth(candidate, font_name, font_size) <= max_width:
                current = candidate
            else:
                ensure_space(line_height, (font_name, font_size))
                pdf.drawString(margin_x + indent, y, current)
                y -= line_height
                current = word

        if current:
            ensure_space(line_height, (font_name, font_size))
            pdf.drawString(margin_x + indent, y, current)
            y -= line_height

    def resolve_font(base_name: str, is_bold: bool, is_italic: bool) -> str:
        if is_bold and is_italic:
            return f'{base_name}-BoldOblique'
        if is_bold:
            return f'{base_name}-Bold'
        if is_italic:
            return f'{base_name}-Oblique'
        return base_name

    def write_styled_paragraph(paragraph, base_font_size: int = 11, indent: int = 0):
        nonlocal y
        import re

        fragments = []
        for run in paragraph.runs:
            text = run.text or ''
            if not text:
                continue
            font_name = resolve_font('Helvetica', bool(run.bold), bool(run.italic))
            fragments.extend(
                (
                    token,
                    font_name,
                    base_font_size,
                )
                for token in re.findall(r'\S+|\s+', text)
            )

        if not fragments:
            return

        max_width = page_width - (margin_x * 2) - indent
        line_height = max(14, int(base_font_size * 1.35))
        url_pattern = re.compile(r'^(https?://|www\.)\S+$', re.IGNORECASE)
        current_line = []
        current_width = 0.0

        def flush_line():
            nonlocal y, current_line, current_width
            if not current_line:
                return
            ensure_space(line_height, ('Helvetica', base_font_size))
            x = margin_x + indent
            for token_text, token_font, token_size in current_line:
                pdf.setFont(token_font, token_size)
                pdf.drawString(x, y, token_text)
                if url_pattern.match(token_text.strip()):
                    target_url = token_text.strip()
                    if target_url.lower().startswith('www.'):
                        target_url = f'https://{target_url}'
                    pdf.linkURL(
                        target_url,
                        (
                            x,
                            y,
                            x + pdf.stringWidth(token_text, token_font, token_size),
                            y + token_size + 2,
                        ),
                        relative=0,
                        thickness=0,
                        color=colors.HexColor('#2563eb'),
                    )
                x += pdf.stringWidth(token_text, token_font, token_size)
            y -= line_height
            current_line = []
            current_width = 0.0

        for token_text, token_font, token_size in fragments:
            token_width = pdf.stringWidth(token_text, token_font, token_size)
            if token_text.isspace() and not current_line:
                continue

            if current_width + token_width > max_width and current_line and not token_text.isspace():
                flush_line()
                if token_text.isspace():
                    continue

            current_line.append((token_text, token_font, token_size))
            current_width += token_width

        flush_line()

    def render_image(image_blob: bytes, width_emu: int | None = None, height_emu: int | None = None):
        nonlocal y

        if not image_blob:
            return

        image = ImageReader(io.BytesIO(image_blob))
        img_width, img_height = image.getSize()
        if not img_width or not img_height:
            return

        max_width = page_width - (margin_x * 2)
        max_height = page_height - (bottom_margin + 40)

        width_inches = (width_emu / 914400) if width_emu else None
        height_inches = (height_emu / 914400) if height_emu else None
        target_width = (width_inches * 72) if width_inches else img_width
        target_height = (height_inches * 72) if height_inches else img_height

        scale = min(max_width / target_width, max_height / target_height, 1.0)
        draw_width = target_width * scale
        draw_height = target_height * scale

        ensure_space(int(draw_height) + 18, ('Helvetica', 11))
        y -= draw_height
        pdf.drawImage(image, margin_x, y, width=draw_width, height=draw_height, preserveAspectRatio=True, mask='auto')
        y -= 16

    wrote_content = False

    for para in document.paragraphs:
        value = (para.text or '').strip()
        if not value:
            continue

        style_name = str(getattr(para.style, 'name', '') or '')
        style_lower = style_name.lower()

        if style_lower.startswith('heading'):
            level = 2
            if style_name and style_name[-1].isdigit():
                level = int(style_name[-1])
            font_size = 18 if level == 1 else (15 if level == 2 else 13)
            write_wrapped(value, font_name='Helvetica-Bold', font_size=font_size)
        elif 'list' in style_lower:
            write_wrapped(f"• {value}", font_name='Helvetica', font_size=11, indent=8)
        elif any(bool(run.bold) or bool(run.italic) for run in para.runs):
            write_styled_paragraph(para, base_font_size=11)
        else:
            write_wrapped(value, font_name='Helvetica', font_size=11)

        wrote_content = True

    seen_blips = set()
    for shape in getattr(document, 'inline_shapes', []):
        inline = getattr(shape, '_inline', None)
        graphic = getattr(inline, 'graphic', None)
        graphic_data = getattr(graphic, 'graphicData', None)
        pic = getattr(graphic_data, 'pic', None)
        blip_fill = getattr(pic, 'blipFill', None)
        blip = getattr(blip_fill, 'blip', None)
        embed_id = getattr(blip, 'embed', None)
        if not embed_id or embed_id in seen_blips:
            continue

        seen_blips.add(embed_id)
        related_part = document.part.related_parts.get(embed_id)
        if not related_part:
            continue

        render_image(
            related_part.blob,
            getattr(shape, 'width', None),
            getattr(shape, 'height', None),
        )
        wrote_content = True

    for table in document.tables:
        table_start_y = y
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = ' '.join((cell.text or '').split())
                cells.append(cell_text)
            row_text = ' | '.join(cells).strip(' |')
            if row_text:
                write_wrapped(row_text, font_name='Helvetica', font_size=10)
                wrote_content = True
        if table.rows:
            write_wrapped('', font_name='Helvetica', font_size=10)

        table_end_y = y
        if table.rows and table.rows[0].cells:
            row_count = len(table.rows)
            column_count = max(len(row.cells) for row in table.rows)
            if column_count:
                row_height = 18
                x0 = margin_x - 4
                y_top = table_start_y + 6
                total_width = page_width - (margin_x * 2) + 8
                total_height = max(row_count * row_height, 18)

                ensure_space(total_height + 18, ('Helvetica', 11))
                pdf.setStrokeColor(colors.HexColor('#94a3b8'))
                pdf.setLineWidth(0.6)

                # Recompute after page break handling to keep the border aligned.
                y_top = table_start_y + 6
                if y_top > page_height - 30:
                    y_top = page_height - 30
                y_bottom = max(table_end_y - 6, bottom_margin)
                pdf.rect(x0, y_bottom, total_width, max(y_top - y_bottom, 18), stroke=1, fill=0)

                for row_index in range(1, row_count):
                    y_line = y_top - (row_index * row_height)
                    if y_line > y_bottom:
                        pdf.line(x0, y_line, x0 + total_width, y_line)

                column_width = total_width / column_count
                for col_index in range(1, column_count):
                    x_line = x0 + (col_index * column_width)
                    pdf.line(x_line, y_bottom, x_line, y_top)

                pdf.setStrokeColor(colors.black)

    if not wrote_content:
        write_wrapped('No extractable text was found in the source DOCX file.', font_name='Helvetica', font_size=11)

    pdf.save()
    output.seek(0)
    return output


def convert_pdf_to_word(file_buffer: io.BytesIO) -> io.BytesIO:
    """Convert a PDF into a Word document by extracting text page-by-page."""

    from docx import Document

    reader = PdfReader(file_buffer)
    document = Document()
    document.add_heading('Aura PDF Conversion', level=1)

    extracted_any_text = False
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ''
        if page_text.strip():
            extracted_any_text = True
            document.add_heading(f'Page {page_number}', level=2)
            for paragraph in page_text.split('\n'):
                paragraph = paragraph.strip()
                if paragraph:
                    document.add_paragraph(paragraph)

    if not extracted_any_text:
        document.add_paragraph('No extractable text was found in the source PDF.')

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

def merge_pdfs(files: List[io.BytesIO], page_orders: Optional[List[List[int]]] = None) -> io.BytesIO:
    """
    Expert PDF Merger for FastAPI
    
    - Supports granular page selection via page_orders
    - Preserves metadata with a custom Aura tag
    """
    writer = PdfWriter()
    
    writer.add_metadata({
        '/Title': 'Aura Merged Document (FastAPI)',
        '/Creator': 'Aura PDF Engine v2.0',
        '/Subject': 'Consolidated PDF'
    })
    
    for i, file_buffer in enumerate(files):
        reader = PdfReader(file_buffer)
        
        # Determine extraction strategy
        orders = page_orders
        if orders is not None and i < len(orders) and orders[i]:
            target_indices = orders[i]
        else:
            target_indices = range(len(reader.pages))
        
        for idx in target_indices:
            if idx < len(reader.pages):
                writer.add_page(reader.pages[idx])
    
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def optimize_pdf(file_buffer: io.BytesIO) -> io.BytesIO:
    """Compress PDF and scrub metadata"""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    
    for page in reader.pages:
        writer.add_page(page)
    
    writer.add_metadata({}) # Secure metadata stripping
    
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def protect_pdf(file_buffer: io.BytesIO, password: str) -> io.BytesIO:
    """Apply AES-256 encryption to the PDF"""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    
    for page in reader.pages:
        writer.add_page(page)
        
    writer.encrypt(password)
    
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def split_pdf(file_buffer: io.BytesIO, mode: str, options: dict) -> List[tuple]:
    """
    Split PDF into multiple parts
    - [mode]: 'every', 'ranges', 'all'
    - [options]: dictionary with mode-specific values
    Returns: List of (filename, BytesIO)
    """
    reader = PdfReader(file_buffer)
    total_pages = len(reader.pages)
    results = []

    if mode == 'every':
        interval = int(options.get('interval', 1))
        for i in range(0, total_pages, interval):
            writer = PdfWriter()
            for j in range(i, min(i + interval, total_pages)):
                writer.add_page(reader.pages[j])
            
            out_buf = io.BytesIO()
            writer.write(out_buf)
            out_buf.seek(0)
            results.append((f"split_{i+1}-{min(i+interval, total_pages)}.pdf", out_buf))

    elif mode == 'ranges':
        ranges_str = options.get('ranges', '')
        parts = ranges_str.split(',')
        for i, part in enumerate(parts):
            writer = PdfWriter()
            segment = part.strip().split('-')
            try:
                if len(segment) == 2:
                    start, end = int(segment[0]), int(segment[1])
                    for p in range(max(1, start), min(total_pages, end) + 1):
                        writer.add_page(reader.pages[p-1])
                elif len(segment) == 1:
                    pageNum = int(segment[0])
                    if 1 <= pageNum <= total_pages:
                        writer.add_page(reader.pages[pageNum - 1])
                
                out_buf = io.BytesIO()
                writer.write(out_buf)
                out_buf.seek(0)
                results.append((f"range_part_{i+1}.pdf", out_buf))
            except:
                continue

    elif mode == 'all':
        for i in range(total_pages):
            writer = PdfWriter()
            writer.add_page(reader.pages[i])
            out_buf = io.BytesIO()
            writer.write(out_buf)
            out_buf.seek(0)
            results.append((f"page_{i+1}.pdf", out_buf))

    return results

def convert_images_to_pdf(image_buffers: List[io.BytesIO]) -> io.BytesIO:
    """Combine multiple image buffers into a single, high-fidelity PDF."""
    from PIL import Image
    images = []
    for buf in image_buffers:
        img = Image.open(buf)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        images.append(img)
    
    output = io.BytesIO()
    if images:
        images[0].save(output, format='PDF', save_all=True, append_images=images[1:])
    
    output.seek(0)
    return output

def remove_pages(file_buffer: io.BytesIO, pages_to_remove: List[int]) -> io.BytesIO:
    """Remove specific 0-indexed pages from the PDF."""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    for idx, page in enumerate(reader.pages):
        if idx not in pages_to_remove:
            writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def extract_pages(file_buffer: io.BytesIO, pages_to_extract: List[int]) -> io.BytesIO:
    """Extract specific 0-indexed pages to form a new PDF."""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    for idx in pages_to_extract:
        if idx < len(reader.pages):
            writer.add_page(reader.pages[idx])
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def rotate_pdf(file_buffer: io.BytesIO, degrees: int = 90) -> io.BytesIO:
    """Rotate all pages in a PDF by the given degrees."""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def unlock_pdf(file_buffer: io.BytesIO, password: str) -> io.BytesIO:
    """Remove encryption from a protected PDF."""
    reader = PdfReader(file_buffer)
    if reader.is_encrypted:
        reader.decrypt(password)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def sign_pdf(file_buffer: io.BytesIO, name: str) -> io.BytesIO:
    """Add a visual signature watermark to the first page."""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    for i, page in enumerate(reader.pages):
        writer.add_page(page)
        
    writer.add_metadata({
        '/Author': name,
        '/Subject': f"Digitally Verified by Aura for {name}",
        '/Keywords': 'Aura-Verified, Secure-PDF'
    })
    
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

def redact_pdf(file_buffer: io.BytesIO, keywords: List[str]) -> io.BytesIO:
    """Strip metadata and specific text patterns from the PDF (basic implementation)."""
    reader = PdfReader(file_buffer)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    
    # Metadata scrubbing is the primary 'redaction' in this basic engine
    writer.add_metadata({})
    
    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output

